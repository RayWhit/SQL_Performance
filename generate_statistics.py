import os
import docx
import matplotlib.pyplot as plt
import numpy as np
import scipy.stats as stats
from scipy.stats import norm, expon
from statistics import mean, stdev, variance
from io import BytesIO


#######################################
#
# Generate overall statistics
#
########################################

def calculate_statistics(file_path):
    with open(file_path, 'r', encoding='latin-1') as file:
        query_times = [float(line.strip()) for line in file]

    mean = np.mean(query_times)
    variance = np.var(query_times)
    std_dev = np.std(query_times)

    return mean, variance, std_dev


def generate_data_dict(folder_paths, statistic='mean'):
    data_dict = {}

    for folder_path in folder_paths:
        folder_data = {}
        for file_name in os.listdir(folder_path):
            if file_name.endswith(".txt"):
                file_path = os.path.join(folder_path, file_name)
                mean, variance, std_dev = calculate_statistics(file_path)

                if statistic == 'mean':
                    folder_data[os.path.splitext(file_name)[0]] = mean
                elif statistic == 'variance':
                    folder_data[os.path.splitext(file_name)[0]] = variance
                elif statistic == 'standard_deviation':
                    folder_data[os.path.splitext(file_name)[0]] = std_dev

        folder_name = os.path.basename(folder_path).replace('queries_times_', '')
        data_dict[folder_name] = folder_data

    # print(data_dict)
    return data_dict


def generate_comparison_chart(folder_paths, statistic='mean'):
    data_dict = generate_data_dict(folder_paths, statistic)

    # Flatten the dictionary to get a list of tuples (folder_name, file_name, stat_value)
    flat_data = [(folder_name, file_name, stat_value) for folder_name, folder_data in data_dict.items() for file_name, stat_value in folder_data.items()]

    # Sort the data alphabetically primarily by file_name and secondarily by folder_name
    sorted_data = sorted(flat_data, key=lambda x: (x[1], x[0]))

    labels = [f"{file_name}" for _, file_name, _ in sorted_data]
    values = [stat_value for _, _, stat_value in sorted_data]
    folder_names = [folder_name for folder_name, _, _ in sorted_data]

    # Sort unique folder names alphabetically for consistency after regenerating
    unique_folder_names = sorted(set(folder_names))

    # Create a color map dictionary based on alphabetical order
    color_map = {folder_name: plt.cm.tab10(i) for i, folder_name in enumerate(unique_folder_names)}
    # Assign colors based on alphabetical order
    colors = [color_map[folder_name] for folder_name in folder_names]

    x = np.arange(len(labels))
    width = 0.3

    fig, ax = plt.subplots(figsize=(10, 5))  # Adjust figure size as needed
    bars = ax.bar(x, values, width, label=statistic.capitalize(), color=colors)

    ax.set_ylabel(f'{statistic.capitalize()} Values')
    ax.set_title(f'Comparison of {statistic.capitalize()}')
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=45, ha='right')

    # Manually create legend handles with appropriate colors
    legend_handles = [plt.Line2D([0], [0], color=color_map[folder_name], lw=4) for folder_name in unique_folder_names]
    legend_labels = [f'{folder}' for folder in unique_folder_names]

    # Add legend with custom handles and labels
    ax.legend(legend_handles, legend_labels)

    # Save the plot to a BytesIO object
    img_buf = BytesIO()
    plt.savefig(img_buf, format='png', dpi=300, bbox_inches='tight', pad_inches=0.3)
    plt.close()

    return img_buf




#######################################
#
# Generate statistics for each DB
#
########################################

def test_normality(data):
    _, p_value = stats.shapiro(data)
    return p_value

def test_exponential(data):
    loc_estimate = min(data)
    scale_estimate = 1 / np.mean(data)
    _, p_value = stats.kstest(data, 'expon', args=(loc_estimate, scale_estimate))
    return p_value

def generate_basic_graph(data, filename):
    plt.figure()
    plt.plot(data, marker='o', linestyle='-')
    plt.title(f"Basic Graph - {filename}")
    plt.xlabel("Query")
    plt.ylabel("Time (ms)")
    plt.grid(True)

    # Save the plot to a BytesIO object
    img_buf = BytesIO()
    plt.savefig(img_buf, format='png')
    plt.close()

    # Return the BytesIO object
    return img_buf

def generate_histogram(data, filename):
    plt.figure()
    
    # Square Root Rule
    num_bins = int(np.sqrt(len(data)))

    # Plot the histogram
    plt.hist(data, bins=num_bins, color='blue', edgecolor='black', density=True, alpha=0.7, label='Data Histogram')

    # Overlay normal distribution curve
    mu, sigma = mean(data), stdev(data)
    x = np.linspace(min(data), max(data), 100)
    y = norm.pdf(x, mu, sigma)
    plt.plot(x, y, 'r-', linewidth=2, label='Normal Distribution')

    # Overlay exponential distribution curve
    loc, scale = min(data), stdev(data)  # Adjusted scale parameter
    y_exponential = expon.pdf(x, loc, scale)
    plt.plot(x, y_exponential, 'g-', linewidth=2, label='Exponential Distribution')
    
    plt.title(f"Histogram - {filename}")
    plt.xlabel("Time (ms)")
    plt.ylabel("Frequency")
    plt.grid(True)

    # Save the plot to a BytesIO object
    img_buf = BytesIO()
    plt.savefig(img_buf, format='png')
    plt.close()

    # Return the BytesIO object
    return img_buf

def generate_stats_docx(folder_paths):
    
    overall_doc = docx.Document()
    
    for stat_type in ['mean', 'variance', 'standard_deviation']:
        # Generate comparison chart and add it to the overall doc
        img_buf = generate_comparison_chart(folder_paths, statistic=stat_type)
        overall_doc.add_heading(f"Comparison Chart - {stat_type.capitalize()}", level=1)
        overall_doc.add_picture(img_buf, width=docx.shared.Inches(7))
        overall_doc.add_page_break()
    
    overall_output_path = os.path.join("./", "overall_stats.docx")
    overall_doc.save(overall_output_path)

    for folder_path in folder_paths:
        doc = docx.Document()
        doc.add_heading(f"Statistics for Files in {folder_path}", level=1)

        for filename in os.listdir(folder_path):
            if filename.endswith(".txt"):
                file_path = os.path.join(folder_path, filename)

                with open(file_path, 'r') as file:
                    content = file.read().splitlines()

                    # Extract numerical data from the file
                    data = [float(line) for line in content if line.strip().replace('.', '').isdigit()]

                    if data:
                        # Basic statistics
                        doc.add_heading(f"File: {filename}", level=2)
                        doc.add_paragraph(f"Number of queries: {round(len(data), 2)}")
                        doc.add_paragraph(f"Mean: {round(np.mean(data), 2)}")
                        doc.add_paragraph(f"Variance: {round(np.var(data), 2)}")
                        doc.add_paragraph(f"Standard Deviation: {round(np.std(data), 2)}")
                        
                        # Perform normality test
                        p_value_normality = test_normality(data)
                        doc.add_paragraph(f"Is normal: {'Yes' if p_value_normality > 0.05 else 'No'}")
                        doc.add_paragraph(f"\tShapiro-Wilk test p-value: {p_value_normality}")

                        # Perform exponential distribution test
                        p_value_exponential = test_exponential(data)
                        doc.add_paragraph(f"Is exponential: {'Yes' if p_value_exponential > 0.05 else 'No'}")
                        doc.add_paragraph(f"\tKolmogorov-Smirnov test p-value: {p_value_exponential}")
                        doc.add_paragraph("\n")

                        # Generate basic graph
                        img_buf_basic = generate_basic_graph(data, filename)
                        doc.add_picture(img_buf_basic, width=docx.shared.Inches(6))

                        # Generate histogram
                        img_buf_histogram = generate_histogram(data, filename)
                        doc.add_picture(img_buf_histogram, width=docx.shared.Inches(6))
                        doc.add_page_break()

        output_docx_path = os.path.join(folder_path, "statistics.docx")
        doc.save(output_docx_path)


if __name__ == "__main__":
    folder_paths = [
        "./queries_times_postgres_single",
        "./queries_times_postgres_multiple",
        "./queries_times_yugabyte",
        "./queries_times_cockroach"
    ]

    # folder_paths = [ "./queries_times_cockroach"]
    
    generate_stats_docx(folder_paths)
    # generate_comparison_chart(folder_paths, statistic='mean')  # For means
    # generate_comparison_chart(folder_paths, statistic='variance')  # For variances
    # generate_comparison_chart(folder_paths, statistic='standard_deviation')  # For standard deviations

